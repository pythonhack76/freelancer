from docx import Document  

import paragraphs

document = Document() 
 
document.add_heading("ciao a tutti", 2 )

p = document.add_paragraph("tutto bene simple text")
p.add_run("new text add").italic = True


document.add_paragraph("questo è un paragrafo")

document.add_paragraph("questo è un paragrafo", style="List Bullet")

document.add_paragraph("questo è un paragrafo",  style="List Bullet")

document.add_paragraph("questo è un paragrafo",  style="List Bullet")


table_header = ["nome", "età", "lavoro"]

some_data = [
    ["pippo", 45, "Programmer"],
      ["mario", 45, "Haredwared"],
        ["enzo", 45, "CEO"],
          ["george", 45, "Programmer"],
            ["ralph", 45, "Programmer"],

]

table = document.add_table(rows=1, cols=3)
for i in range(3):
    table.rows[0].cells[i].text = table_header[i]

for nome, anni, lavoro in some_data:
    cells = table.add_row().cells
    cells[0].text = nome
    cells[1].text = str(anni)  
    cells[2].text = lavoro  

document.save("test.docx")

